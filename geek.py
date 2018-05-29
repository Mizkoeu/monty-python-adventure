import re
from docx.shared import Pt
from docx import Document
from googlesearch import search
from bs4 import BeautifulSoup as soup
from requests import get
from requests.exceptions import RequestException
from contextlib import closing

'A program to fetch article data from GeeksForGeeks website and convert it into
a docx/pdf document for easier reading'

def simple_get(url):
    try:
        with closing(get(url, stream=True)) as resp:
            if is_good_response(resp):
                return resp.content
            else:
                return None

    except RequestException as e:
        print('Error during requests to {0} : {1}'.format(url, str(e)))
        return None

def is_good_response(resp):
    content_type = resp.headers['Content-Type'].lower()
    return (resp.status_code == 200 
            and content_type is not None 
            and content_type.find('html') > -1)


query_str = input('What topic would you like to learn about? ')
num_page = input('How many pages do you want me to fetch? ')

document = Document()
document.add_heading(query_str.upper(), 0)

query_str += ' geeksforgeeks'
for raw_url in search(query_str, tld="co.in", num=int(num_page), stop=1, pause=2):
    raw_html = simple_get(raw_url)
    html = soup(raw_html, 'html.parser')
    # parse the page for content
    # content = html.find('div', {'class':'entry-content'})
    content = html.select_one('div.entry-content')
    if content != None:
        n=1 
        for p in content.find_all('pre'):
            print(p.text)
            p.extract()
        for p in content.find_all('div'):
            p.extract()
        for p in content.find_all('p'): 
            if p.text == '': 
                continue
            print(str(n) + '. ' + p.name)
            if p.name == 'pre': 
                print('This is a PREFORMATTED STRING!!!') 
                document.add_paragraph(str(n) + p.text)
                n+=1
            else:
                document.add_paragraph(str(n) + p.text)
                # print(str(n) + '. ' + p.name)
                n+=1

document.save(query_str.replace(' ', '_') + '.docx')


print("-----------------" + query_str + "------------------------------\n")

# raw_html = simple_get('http://marvel.wikia.com/wiki/Marvel_Database')
# html = soup(raw_html, 'html.parser')
# for p in html.find_all('a', href=True, title=True):
#     if p['href'].startswith('/wiki/'):
#         print(p['title'])
